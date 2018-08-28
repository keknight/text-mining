
import sys
import os
import itertools
import nltk
import docx
import pandas as pd
from gensim.models import Phrases
from gensim.models.phrases import Phraser
from gensim.parsing.preprocessing import remove_stopwords, strip_punctuation

nltk.download('averaged_perceptron_tagger')


#function to read text inside word doc tables
def readTables(doc):
	x = []
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					x.append(paragraph.text)
	return x

	
#function to identify tagged nouns in corpus	
is_noun = lambda pos: pos[:2] == 'NN'


#filename = sys.argv[-1]
#doc = docx.Document(filename)
#tables = readTables(doc)

directory = 'C:/Users/vkk/Desktop/job posts'
extension = '.docx'


'''
files = []
for file in os.listdir(directory):
	if file.endswith(extension):
		files.append(readTables(docx.Document(file)))
'''

files = [readTables(docx.Document(file)) for file in os.listdir(directory) if file.endswith(extension)]
 
#read docs, convert each doc's tables to list of lists
#length of outer list is number of docs in file path
#remove blanks, flatten each doc into one long list of words 

texts = [[[word for word in table.split()] for table in file] for file in files]
flat_texts = [list(itertools.chain.from_iterable(x)) for x in texts]
flat_texts = [[strip_punctuation(remove_stopwords(word.lower())) for word in group] for group in flat_texts]
texts = [[item for item in table if item != ''] for table in flat_texts]

d = {'CCSD': texts[0], 'EESD': texts[1], 'GSD': texts[2], 'NScD': texts[3], 'NSED': texts[4], 'PSD': texts[5]}

df = pd.DataFrame(dict([(k, pd.Series(v)) for k, v in d.items()]))

#find intersection of series
#TODO: figure out a way to do this automatically

compared = pd.Series(list(set(df['EESD']) & set(df['GSD'])))
compared2 = pd.Series(list(set(df['EESD']) & set(df['NSED'])))
compared3 = pd.Series(list(set(df['EESD']) & set(df['NScD'])))
compared4 = pd.Series(list(set(df['EESD']) & set(df['PSD'])))
compared5 = pd.Series(list(set(df['EESD']) & set(df['CCSD'])))

eesd = pd.DataFrame(dict(gsd = compared, nsed = compared2, nscd = compared3, psd = compared4, ccsd = compared5)).reset_index()


writer = pd.ExcelWriter('eesd_compared.xlsx', engine = 'xlsxwriter')
eesd.to_excel(writer, encoding = 'ISO-8859-1')
writer.save()

'''
phrases = Phrases(texts)
bigram = Phraser(phrases)
trigram = Phrases(bigram[texts])
trigram = Phraser(trigram)
phTexts = [trigram[bigram[text]] for text in texts]
'''

#create list of just nouns
nouns = [[word for (word, pos) in nltk.pos_tag(text) if is_noun(pos)] for text in flat_texts]

################TO DO#####################

#create dataframe of all texts, then compare each column using intersection
set(nouns).intersection(other_nouns)
#compares first item in list...
result = set(d[0]).intersection(*d[:1])
#not sure yet
set.intersection(*map(set,d))

#or try this
[x for x in nouns if x in other_nouns]


'''
df = nouns.to_frame(name = filename)
#figure out how to read other files, add data to additional columns in dataframe
df['new_col'] = se.values

##convert files to dataframe columns, compare columns using python diff
#do one at a time? column[0] to column[x], export results, then column[1] to column[x]
'''